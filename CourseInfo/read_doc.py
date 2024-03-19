from docx import Document
import re

class CourseInfo():
    def __init__(self, path):
        self._file_path = path

    @property
    def file_path(self):
        return self._file_path
   
    @property
    def document(self):
        return self._doc



    def read_ig(self):
        self._doc = Document(self._file_path)
    
    def course_id(self, string):
        numbers = re.findall(r'\d+', string)
        self._course_id = [num for num in numbers][0]
    
    def get_ptbs(self):
        self.ptb = {}
        
        location = [num for num,x in enumerate(self._doc.paragraphs) if x.text.upper().startswith("PTB TASKS COVERED")]
        pattern = r"^\d+|PTB\d+"
        if location:
            num = location[0]
            num+1
            
            p = self._doc.paragraphs[num]

            while num < len(self._doc.paragraphs) and not p.text.lower().startswith('knowledge') and not p.text.lower().startswith("not applicable"):
                match = re.match(pattern, p.text.strip()) #check for match
                
                if match: 
                    self.ptb[re.sub(r'\D', '', match.group())] = p.text[match.end():].replace(":", "").strip()
                num+=1
                p = self._doc.paragraphs[num]

    def get_dhs(self):
        self.dhs = []
        
        location = [num for num,x in enumerate(self._doc.paragraphs) if x.text.upper().startswith("DHS COMPETENCIES COVERED")]

        if location:
            
            for num in location:
                num+=1

                p = self._doc.paragraphs[num]

                while num < len(self._doc.paragraphs) and (not p.text.lower().startswith('ptb tasks') and not p.text.lower().startswith("not applicable") and not p.text.lower().startswith('knowledge') and not p.text.lower().startswith('note:')):
                    self.dhs.append(p.text.strip())
                    num+=1
                    p = self._doc.paragraphs[num]

    def get_tlo(self):
        self.tlo = []
        location = [num for num,x in enumerate(self._doc.paragraphs) if x.text.upper().startswith("TERMINAL LEARNING")]


        if location:
            num = location[0]
            num+=1

            p = self._doc.paragraphs[num]

            while num < len(self._doc.paragraphs) and not p.text.lower().startswith('enabling learning') and not  p.text.lower().startswith("not applicable") and not p.text.lower().startswith("note:") and not p.text.lower().startswith('knowledge') and not p.text.lower().startswith('ptb tasks'):
                self.tlo.append(p.text)
                num+=1
                p = self._doc.paragraphs[num]

        # print(self.tlo)
    

    def get_elos(self):
        self.elos = []
        location = [num for num,x in enumerate(self._doc.paragraphs) if x.text.upper().startswith("ENABLING LEARNING") or x.text.upper().startswith("ELO")]
        if location:
            
            for num in location:
                num+=1

                p = self._doc.paragraphs[num]

                while num < len(self._doc.paragraphs) and not p.text.lower().startswith('dhs competenc') and not p.text.lower().startswith('ptb tasks') and not p.text.lower().startswith("not applicable") and not p.text.lower().startswith('knowledge') and not p.text.lower().startswith('note:'):
                    self.elos.append(p.text.strip())
                    num+=1
                    p = self._doc.paragraphs[num]

        # print(self.elos)

if __name__ == "__main__":
    test = CourseInfo(r'G:\.shortcut-targets-by-id\1LHJ1O_oPDPAnLMGwF4cc8OdBxGumjxpQ\Training\IA\ASSP\Final\IGs\FQS-IA_0836_ASSP_Module1_IG_03-22-24.docx')
    test.read_ig()
    test.get_ptbs()
    print(test.ptb)
    print(test.ptb_text)